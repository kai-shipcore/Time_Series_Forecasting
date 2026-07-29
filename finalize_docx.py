#!/usr/bin/env python3
"""Post-process the pandoc-built proposal docx.

Run by build_docx.sh after pandoc. It:
  1. Gives every table a solid gridline border and a shaded header row.
  2. Keeps each table on a single page (no splitting across a page break) and
     keeps it with its caption.
  3. Tightens the spacing between bullet points.

Needs python-docx (pip install python-docx).
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BORDER_COLOR = "BFBFBF"
HEADER_FILL = "EDEDED"


def _el(tag):
    return OxmlElement(tag)


def table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    for b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(b)
    borders = _el("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = _el("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), BORDER_COLOR)
        borders.append(e)
    look = tblPr.find(qn("w:tblLook"))
    if look is not None:
        look.addprevious(borders)
    else:
        tblPr.append(borders)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = _el("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    valign = tcPr.find(qn("w:vAlign"))
    if valign is not None:
        valign.addprevious(shd)
    else:
        tcPr.append(shd)


def cant_split(row):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = _el("w:trPr")
        tr.insert(0, trPr)
    trPr.append(_el("w:cantSplit"))


def keep_next(p_elem):
    pPr = p_elem.get_or_add_pPr()
    if pPr.find(qn("w:keepNext")) is not None:
        return
    kn = _el("w:keepNext")
    pstyle = pPr.find(qn("w:pStyle"))
    if pstyle is not None:
        pstyle.addnext(kn)
    else:
        pPr.insert(0, kn)


def main(path):
    d = Document(path)

    for tbl in d.tables:
        table_borders(tbl)
        for cell in tbl.rows[0].cells:
            shade_cell(cell, HEADER_FILL)
        rows = tbl.rows
        for i, row in enumerate(rows):
            cant_split(row)
            if i < len(rows) - 1:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        keep_next(p._p)

    # keep each table with the caption paragraph immediately above it
    body = d.element.body
    kids = list(body)
    for i, ch in enumerate(kids):
        if ch.tag == qn("w:tbl") and i > 0 and kids[i - 1].tag == qn("w:p"):
            keep_next(kids[i - 1])

    # tighten spacing on bullet/numbered list items
    for p in d.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

    d.save(path)
    print("finalized", path)


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "Machine_Learning_Demand_Forecast_Proposal.docx")
